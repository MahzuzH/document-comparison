import streamlit as st
import os
from PyPDF2 import PdfReader
from PyPDF2 import PdfFileReader
import PyPDF2
import docx
import pandas as pd
from io import StringIO
import string
import re
import nltk
from nltk.corpus import stopwords
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
from Sastrawi.StopWordRemover.StopWordRemoverFactory import StopWordRemoverFactory

def stemming(text):
    #penghilangan white space
    whitespace = text.translate(str.maketrans('','',string.punctuation)).replace("\t", "").replace("\n", "").replace("\r", "").strip()
    #mengubah menjadi lowercase
    lowercase = whitespace.lower()
    #menghapus angka
    removed_angka = re.sub(r"\d+", "", lowercase)
    hasil = removed_angka
    st.text_area("Hasil Preprocessing: ", hasil, disabled=True)
    
    #melakukan tokenizing, yaitu memisahkan kata perkata
    tokenized = nltk.tokenize.word_tokenize(hasil)
    st.text_area("Hasil Tokenizing: ", tokenized, disabled=True)
    
    # fungsi stopwording/filtering yaitu menghilangkan kata-kata yang tidak penting
    #list stopword bahasa indonesia dari library nltk
    listStopword = nltk.corpus.stopwords.words('indonesian')
    #penambahan kata stopword
    addstop = ['ku', 'pel', 'uru', 'mem', 'astikan', 's', 'iapa', 'koyak']
    listStopword.extend(addstop)
    #proses filtering
    filtering = [words for words in tokenized if not words in listStopword]    
    st.text_area("Hasil Stopword: ", filtering, disabled=True) 
    
    # stemming mengubah kata menjadi kata dasarnya
    stemmer = StemmerFactory().create_stemmer()
    sentence_doc = ' '.join(map(str, filtering))
    stemming_doc = stemmer.stem(sentence_doc)
    result_doc = nltk.tokenize.word_tokenize(stemming_doc)
    st.text_area("Hasil Stemming: ", result_doc, disabled=True)
    return result_doc


st.title("IR Apps - Mencari data dari query")
st.subheader("Kelompok 2")

data = []
text =[]
# Get the directory path from the user
try:
  dir_path = st.text_input("Masukkan path folder: ")
  # Get a list of all the files in the directory
  files = os.listdir(dir_path)
  os.chdir(dir_path)
  
  
  # Iterate over the list of files
  for file in files:
    file_path = os.path.join(dir_path, file)  # construct the full path to the file
    file_details = os.stat(file_path)  # get the file's details
    file_name, file_ext = os.path.splitext(file)
    data.append((file_name, file_ext))
    
    
  # table_data = {'Nama File': file, 'Size': file_details.st_size, 'Last modified:': file_details.st_mtime, 'Extention: ': file_extention[1]}
  df = pd.DataFrame(data, columns=['Nama File', 'Ekstensi File'])
  st.write(df)
  
except:
  st.write(" ")
 
 
if st.button("Stemming"):
  for i, ekstrak in enumerate(files):
    #mengambil ektensi
    file_name, file_ext = os.path.splitext(ekstrak)
    if file_ext == '.docx':
      doc = docx.Document(ekstrak)
      fulldoc = []
      # all_parag_doc = doc.paragraphs
      st.subheader("Hasil ekstrak dokumen ke-"+ f"{i}" + " (DOCX)")
      for parag in doc.paragraphs:
          fulldoc.append(parag.text)
          
      test =  ' '.join(map(str, fulldoc))    
      st.write(test)
      hasilstem_docx = stemming(test)
        
    elif file_ext == '.pdf':
      with open(ekstrak, 'rb') as file: 
        # Baca file PDF
        reader = PyPDF2.PdfFileReader(file)
        # Ambil jumlah halaman
        jumlah_halaman = reader.getNumPages()

        # Iterasi melalui setiap halaman
        for halaman in range(jumlah_halaman):
            # Baca halaman ke-i
            hal = reader.getPage(halaman)
            # Dapatkan teks dari halaman
            teks = hal.extractText()
      st.subheader("Hasil ekstrak dokumen ke-"+f"{i}"+ " (PDF)")
      st.write(teks)
      hasilstem_pdf = stemming(teks)
      
    else:
      st.write("Tipe file salah, hanya file pdf dan docx yang dapat dilakukan stemming")
            